#!/usr/bin/env python3
"""
Word MCP Server - Universal Microsoft Word Document Manipulation
Embedded in PowerNode as child process via MCP Python SDK
"""

import asyncio
import os
import sys
import json
from io import BytesIO
from typing import Any, Dict, List, Optional

# MCP SDK imports
try:
    from mcp.server import Server
    from mcp.server.stdio import stdio_server
    from mcp import types
except ImportError:
    print("ERROR: MCP SDK not installed. Run: pip install mcp", file=sys.stderr)
    sys.exit(1)

# Word processing imports
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Azure Storage
from azure.storage.blob import BlobServiceClient

# Anthropic AI
from anthropic import Anthropic

# Initialize server
app = Server("word-mcp-server")

# Global clients (initialized from environment)
blob_service_client: Optional[BlobServiceClient] = None
anthropic_client: Optional[Anthropic] = None
default_container: str = "wippli-documents"

def init_clients():
    """Initialize Azure and Anthropic clients from environment variables"""
    global blob_service_client, anthropic_client, default_container

    # Azure Blob Storage
    azure_conn = os.getenv("AZURE_STORAGE_CONNECTION_STRING")
    if azure_conn:
        blob_service_client = BlobServiceClient.from_connection_string(azure_conn)

    # Anthropic API
    anthropic_key = os.getenv("ANTHROPIC_API_KEY")
    if anthropic_key:
        anthropic_client = Anthropic(api_key=anthropic_key)

    # Default container
    default_container = os.getenv("DEFAULT_CONTAINER", "wippli-documents")

# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

async def download_doc_from_blob(file_path: str, container: Optional[str] = None) -> Document:
    """Download Word document from Azure Blob Storage"""
    if not blob_service_client:
        raise ValueError("Azure Blob Storage not configured")

    container_name = container or default_container
    blob_client = blob_service_client.get_blob_client(container_name, file_path)

    doc_bytes = blob_client.download_blob().readall()
    return Document(BytesIO(doc_bytes))

async def upload_doc_to_blob(doc: Document, file_path: str, container: Optional[str] = None):
    """Upload Word document to Azure Blob Storage"""
    if not blob_service_client:
        raise ValueError("Azure Blob Storage not configured")

    container_name = container or default_container
    blob_client = blob_service_client.get_blob_client(container_name, file_path)

    # Save to BytesIO
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    # Upload
    blob_client.upload_blob(doc_io, overwrite=True)

# ============================================================================
# MCP TOOLS
# ============================================================================

@app.list_tools()
async def list_tools() -> List[types.Tool]:
    """List all available Word MCP tools"""
    return [
        types.Tool(
            name="create_document",
            description="Create a new Word document from scratch with optional title and initial content",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "Path where document will be saved in blob storage"},
                    "title": {"type": "string", "description": "Document title"},
                    "content": {"type": "string", "description": "Initial paragraph content"},
                    "container": {"type": "string", "description": "Azure blob container (optional)"},
                },
                "required": ["file_path"],
            },
        ),
        types.Tool(
            name="read_document",
            description="Read and extract all content from a Word document including structure, text, tables, and metadata",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "Path to document in blob storage"},
                    "container": {"type": "string", "description": "Azure blob container (optional)"},
                },
                "required": ["file_path"],
            },
        ),
        types.Tool(
            name="list_documents",
            description="List all Word documents (.docx) in Azure blob container with metadata",
            inputSchema={
                "type": "object",
                "properties": {
                    "container": {"type": "string", "description": "Azure blob container (optional)"},
                    "prefix": {"type": "string", "description": "Filter by path prefix"},
                },
            },
        ),
        types.Tool(
            name="add_paragraph",
            description="Add a paragraph to an existing Word document",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "Path to document in blob storage"},
                    "text": {"type": "string", "description": "Paragraph text to add"},
                    "container": {"type": "string", "description": "Azure blob container (optional)"},
                },
                "required": ["file_path", "text"],
            },
        ),
        types.Tool(
            name="analyze_questionnaire",
            description="AI-powered analysis of questionnaire structure - detects sections, questions, form fields, rating scales automatically (works with ANY questionnaire in ANY language)",
            inputSchema={
                "type": "object",
                "properties": {
                    "file_path": {"type": "string", "description": "Path to questionnaire document"},
                    "container": {"type": "string", "description": "Azure blob container (optional)"},
                },
                "required": ["file_path"],
            },
        ),
    ]

@app.call_tool()
async def call_tool(name: str, arguments: Any) -> List[types.TextContent]:
    """Handle tool execution"""

    try:
        if name == "create_document":
            return await create_document_tool(arguments)
        elif name == "read_document":
            return await read_document_tool(arguments)
        elif name == "list_documents":
            return await list_documents_tool(arguments)
        elif name == "add_paragraph":
            return await add_paragraph_tool(arguments)
        elif name == "analyze_questionnaire":
            return await analyze_questionnaire_tool(arguments)
        else:
            return [types.TextContent(type="text", text=f"Unknown tool: {name}")]

    except Exception as e:
        error = {"success": False, "error": str(e)}
        return [types.TextContent(type="text", text=json.dumps(error, indent=2))]

# ============================================================================
# TOOL IMPLEMENTATIONS
# ============================================================================

async def create_document_tool(args: Dict[str, Any]) -> List[types.TextContent]:
    """Create new Word document"""
    doc = Document()

    # Add title if provided
    if args.get("title"):
        doc.add_heading(args["title"], 0)

    # Add initial content if provided
    if args.get("content"):
        doc.add_paragraph(args["content"])

    # Upload to blob
    await upload_doc_to_blob(doc, args["file_path"], args.get("container"))

    result = {
        "success": True,
        "file_path": args["file_path"],
        "container": args.get("container") or default_container,
        "message": "Document created successfully"
    }

    return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

async def read_document_tool(args: Dict[str, Any]) -> List[types.TextContent]:
    """Read document structure and content"""
    doc = await download_doc_from_blob(args["file_path"], args.get("container"))

    # Extract structure
    result = {
        "file_path": args["file_path"],
        "paragraphs": [
            {"text": p.text, "style": p.style.name}
            for p in doc.paragraphs if p.text.strip()
        ],
        "tables": [
            {
                "rows": len(table.rows),
                "columns": len(table.columns) if table.rows else 0,
                "data": [
                    [cell.text for cell in row.cells]
                    for row in table.rows
                ]
            }
            for table in doc.tables
        ],
        "sections": len(doc.sections),
    }

    return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

async def list_documents_tool(args: Dict[str, Any]) -> List[types.TextContent]:
    """List all Word documents in blob container"""
    if not blob_service_client:
        raise ValueError("Azure Blob Storage not configured")

    container_name = args.get("container") or default_container
    container_client = blob_service_client.get_container_client(container_name)

    # List blobs
    prefix = args.get("prefix", "")
    blobs = container_client.list_blobs(name_starts_with=prefix)

    documents = [
        {
            "name": blob.name,
            "size": blob.size,
            "last_modified": blob.last_modified.isoformat() if blob.last_modified else None,
        }
        for blob in blobs
        if blob.name.endswith(".docx")
    ]

    result = {
        "container": container_name,
        "count": len(documents),
        "documents": documents,
    }

    return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

async def add_paragraph_tool(args: Dict[str, Any]) -> List[types.TextContent]:
    """Add paragraph to existing document"""
    doc = await download_doc_from_blob(args["file_path"], args.get("container"))
    doc.add_paragraph(args["text"])
    await upload_doc_to_blob(doc, args["file_path"], args.get("container"))

    result = {
        "success": True,
        "message": "Paragraph added successfully",
        "file_path": args["file_path"]
    }

    return [types.TextContent(type="text", text=json.dumps(result, indent=2))]

async def analyze_questionnaire_tool(args: Dict[str, Any]) -> List[types.TextContent]:
    """AI-powered questionnaire structure analysis"""
    if not anthropic_client:
        raise ValueError("Anthropic API not configured")

    # Download and parse document
    doc = await download_doc_from_blob(args["file_path"], args.get("container"))

    # Extract document structure for AI analysis
    structure = {
        "paragraphs": [
            {"index": i, "text": p.text, "style": p.style.name}
            for i, p in enumerate(doc.paragraphs) if p.text.strip()
        ],
        "tables": [
            {
                "index": i,
                "rows": len(table.rows),
                "columns": len(table.columns) if table.rows else 0,
                "sample_data": [
                    [cell.text for cell in row.cells]
                    for row in list(table.rows)[:3]  # First 3 rows as sample
                ]
            }
            for i, table in enumerate(doc.tables)
        ],
    }

    # Use Claude to understand structure
    prompt = f"""Analyze this Word document structure and identify all questionnaire elements:

Document structure:
{json.dumps(structure, indent=2)}

Please identify:
1. All questions and their types (text input, multiple choice, rating scale, yes/no, etc.)
2. Section organization
3. Required vs optional fields
4. Rating scales (if any) with their options
5. Table-based questions

Return a JSON object with this structure:
{{
  "document_type": "questionnaire",
  "total_questions": <number>,
  "sections": [
    {{
      "title": "<section name>",
      "questions": [
        {{
          "id": "<unique_id>",
          "text": "<question text>",
          "type": "<question_type>",
          "location": "<paragraph_index or table_index>",
          "options": ["<option1>", "<option2>"] (if applicable),
          "required": true/false
        }}
      ]
    }}
  ]
}}"""

    message = anthropic_client.messages.create(
        model="claude-3-5-sonnet-20241022",
        max_tokens=4096,
        messages=[{"role": "user", "content": prompt}]
    )

    # Parse AI response
    analysis = json.loads(message.content[0].text)

    return [types.TextContent(type="text", text=json.dumps(analysis, indent=2))]

# ============================================================================
# MAIN ENTRY POINT
# ============================================================================

async def main():
    """Main entry point for MCP server"""
    # Initialize clients
    init_clients()

    # Run server (stdio transport for child process communication)
    async with stdio_server() as (read_stream, write_stream):
        await app.run(
            read_stream,
            write_stream,
            app.create_initialization_options()
        )

if __name__ == "__main__":
    asyncio.run(main())
